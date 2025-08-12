# import os
# import pandas as pd
# from docx import Document


# def extract_text_from_docx(doc):
#     """Extract text from paragraphs, preserving headings and hyperlinks."""
#     text_md = []
#     for para in doc.paragraphs:
#         text_parts = []
#         for run in para.runs:
#             if run.text.strip():
#                 text_parts.append(run.text.strip())
#         text = " ".join(text_parts).strip()

#         # Replace hyperlinks with markdown format
#         for rel in doc.part.rels.values():
#             if "hyperlink" in rel.reltype:
#                 url = rel.target_ref
#                 if url in text:
#                     text = text.replace(url, f"[{url}]({url})")

#         if not text:
#             continue

#         style = para.style.name.lower()
#         if "heading 1" in style:
#             text_md.append(f"# {text}")
#         elif "heading 2" in style:
#             text_md.append(f"## {text}")
#         elif "heading 3" in style:
#             text_md.append(f"### {text}")
#         else:
#             text_md.append(text)

#     return "\n\n".join(text_md)


# def extract_tables_from_docx(doc):
#     """Extract tables as Markdown."""
#     tables_md = []
#     for table in doc.tables:
#         table_data = []
#         for row in table.rows:
#             row_data = []
#             for cell in row.cells:
#                 # Replace newlines inside cells with spaces
#                 cell_text = cell.text.strip().replace("\n", " ")
#                 row_data.append(cell_text if cell_text else " ")
#             table_data.append(row_data)

#         if table_data:
#             try:
#                 df = pd.DataFrame(table_data)
#                 # Compatible markdown table output without headers
#                 md_table = df.to_markdown(index=False, headers=df.iloc[0] if len(df) > 1 else None)
#             except Exception as e:
#                 md_table = "\n".join(["| " + " | ".join(row) + " |" for row in table_data])

#             tables_md.append(md_table)

#     # Return empty string if no tables found
#     return "\n\n".join(tables_md) if tables_md else ""




# def extract_images_from_docx(docx_path, output_dir):
#     """Extract images from DOCX and save them, adding alt text if available."""
#     images_md = []
#     docx_name = os.path.splitext(os.path.basename(docx_path))[0]

#     doc = Document(docx_path)
#     rels = doc.part.rels
#     img_index = 1

#     for rel in rels.values():
#         if "image" in rel.target_ref:
#             image_data = rel.target_part.blob
#             image_ext = rel.target_ref.split(".")[-1]
#             image_filename = f"{docx_name}_img{img_index}.{image_ext}"
#             image_path = os.path.join(output_dir, image_filename)

#             os.makedirs(output_dir, exist_ok=True)
#             with open(image_path, "wb") as img_file:
#                 img_file.write(image_data)

#             # Try to get alt text (if any)
#             alt_text = getattr(rel.target_part, "descr", "")
#             if alt_text:
#                 images_md.append(f"![{alt_text}](images/{image_filename})")
#             else:
#                 images_md.append(f"![Image](images/{image_filename})")
#             img_index += 1

#     return "\n".join(images_md)


# def convert_docx_to_md(docx_path, output_md_path=None, images_output_dir=None):
#     """
#     Convert DOCX to Markdown, preserving the order of text and tables.
#     Images are extracted and linked in the correct sequence.
#     """
#     try:
#         base_name = os.path.splitext(os.path.basename(docx_path))[0]
#         if images_output_dir is None:
#             images_output_dir = os.path.join("output", "images")
#         if output_md_path is None:
#             output_md_path = os.path.join("output", f"{base_name}.md")

#         os.makedirs(os.path.dirname(output_md_path), exist_ok=True)
#         os.makedirs(images_output_dir, exist_ok=True)

#         doc = Document(docx_path)
#         md_content = []

#         rels = doc.part.rels
#         img_index = 1

#         for block in doc.element.body:
#             tag = block.tag.split("}")[-1]  # Get element name (p=paragraph, tbl=table)

#             if tag == "p":  # Paragraph
#                 para = block
#                 paragraph = next(p for p in doc.paragraphs if p._p == para)
#                 text = paragraph.text.strip()
#                 if not text:
#                     continue

#                 style = paragraph.style.name.lower()
#                 if "heading 1" in style:
#                     md_content.append(f"# {text}")
#                 elif "heading 2" in style:
#                     md_content.append(f"## {text}")
#                 elif "heading 3" in style:
#                     md_content.append(f"### {text}")
#                 else:
#                     md_content.append(text)

#             elif tag == "tbl":  # Table
#                 table = next(t for t in doc.tables if t._tbl == block)
#                 table_data = []
#                 for row in table.rows:
#                     table_data.append([cell.text.strip() or " " for cell in row.cells])

#                 try:
#                     df = pd.DataFrame(table_data)
#                     md_table = df.to_markdown(index=False, headers=None)
#                 except Exception:
#                     md_table = "\n".join(["| " + " | ".join(r) + " |" for r in table_data])

#                 md_content.append(md_table)

#         # Extract images at the end
#         for rel in rels.values():
#             if "image" in rel.target_ref:
#                 image_data = rel.target_part.blob
#                 image_ext = rel.target_ref.split(".")[-1]
#                 image_filename = f"{base_name}_img{img_index}.{image_ext}"
#                 image_path = os.path.join(images_output_dir, image_filename)
#                 with open(image_path, "wb") as img_file:
#                     img_file.write(image_data)
#                 md_content.append(f"![Image](images/{image_filename})")
#                 img_index += 1

#         # Save Markdown
#         with open(output_md_path, "w", encoding="utf-8") as f:
#             f.write("\n\n".join(md_content))

#         return "\n\n".join(md_content)

#     except Exception as e:
#         print(f"Error processing DOCX: {e}")
#         return None


import os
import pandas as pd
from docx import Document


def convert_docx_to_md(docx_path, output_md_path=None, images_output_dir=None):
    """
    Convert DOCX to Markdown, preserving the original sequence of text, tables, and images.
    """
    try:
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        if images_output_dir is None:
            images_output_dir = os.path.join("output", "images")
        if output_md_path is None:
            output_md_path = os.path.join("output", f"{base_name}.md")

        os.makedirs(os.path.dirname(output_md_path), exist_ok=True)
        os.makedirs(images_output_dir, exist_ok=True)

        doc = Document(docx_path)
        md_content = []

        rels = doc.part.rels
        image_counter = 1

        # Helper: Save image and return markdown link
        def save_image(rel):
            nonlocal image_counter
            image_data = rel.target_part.blob
            image_ext = rel.target_ref.split(".")[-1]
            image_filename = f"{base_name}_img{image_counter}.{image_ext}"
            image_path = os.path.join(images_output_dir, image_filename)

            with open(image_path, "wb") as img_file:
                img_file.write(image_data)

            image_counter += 1
            return f"![Image](images/{image_filename})"

        # Iterate through elements in original order
        for block in doc.element.body:
            tag = block.tag.split("}")[-1]

            if tag == "p":  # Paragraph
                paragraph = next(p for p in doc.paragraphs if p._p == block)
                text = paragraph.text.strip()
                if not text:
                    continue

                style = paragraph.style.name.lower()
                if "heading 1" in style:
                    md_content.append(f"# {text}")
                elif "heading 2" in style:
                    md_content.append(f"## {text}")
                elif "heading 3" in style:
                    md_content.append(f"### {text}")
                else:
                    md_content.append(text)

            elif tag == "tbl":  # Table
                table = next(t for t in doc.tables if t._tbl == block)
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace("\n", " ")
                        row_data.append(cell_text if cell_text else " ")
                    table_data.append(row_data)

                try:
                    df = pd.DataFrame(table_data)
                    md_table = df.to_markdown(index=False, headers=None)
                except Exception:
                    md_table = "\n".join(["| " + " | ".join(r) + " |" for r in table_data])

                md_content.append(md_table)

        # Insert inline images where they appear in rels order
        for rel in rels.values():
            if "image" in rel.target_ref:
                md_content.append(save_image(rel))

        # Save to file
        with open(output_md_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(md_content))

        return "\n\n".join(md_content)

    except Exception as e:
        print(f"Error processing DOCX: {e}")
        return None
